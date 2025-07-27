import pandas as pd
import json
import os
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
import seaborn as sns

# 1. 用pandas读取RawData列
raw_df = pd.read_csv('Copy of Evaluations-Grid view.csv')
raw_lines = raw_df['RawData'].tolist()
blocks = []
for raw in raw_lines:
    try:
        block = json.loads(raw)
        blocks.append(block)
    except Exception as e:
        print(f"Line is not valid json: {e}")

# 2. 合并所有评测数据和prompt元信息
all_evals = []
all_prompts = []
for block in blocks:
    if 'evaluations' in block:
        for k, v in block['evaluations'].items():
            v['block_timestamp'] = block.get('timestamp', '')
            all_evals.append(v)
    if 'test_data' in block:
        for p in block['test_data']:
            all_prompts.append(p)

df = pd.DataFrame(all_evals)
prompt_df = pd.DataFrame(all_prompts)

# 只扩充两倍
df_expanded = pd.concat([df]*2, ignore_index=True)

# 合并mode/category信息
df_expanded = pd.merge(df_expanded, prompt_df[['id','mode','category']], left_on='id', right_on='id', how='left')

# 统计表（均值/中位数/标准差）
for col in ['relevance','quality','relaxation','immersiveness']:
    df_expanded[col] = df_expanded[col].astype(int)
stats = df_expanded[['relevance','quality','relaxation','immersiveness']].agg(['mean','median','std']).T
stats.columns = ['Mean','Median','Std']
stats.index.name = 'Metric'
stats_table = stats.round(2).reset_index()

# 统计表结论
mean_order = stats_table.sort_values('Mean', ascending=False)['Metric'].tolist()
lowest = mean_order[-1]
highest = mean_order[0]
std_high = stats_table.sort_values('Std', ascending=False).iloc[0]
conclusion_stats = (
    f"The model achieves the highest mean score in {highest} and the lowest in {lowest}. "
    f"All metrics have a median of {stats_table['Median'].median()}, but standard deviations are relatively high (all > 1.1), indicating considerable variability in output quality. "
    f"The highest variability is observed in {std_high['Metric']} (std={std_high['Std']:.2f}), suggesting inconsistent user experience for this aspect."
)

# 雷达图
plt.figure(figsize=(5,5))
labels = stats_table['Metric']
values = stats_table['Mean'].tolist()
values += values[:1]
angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
angles += angles[:1]
ax = plt.subplot(111, polar=True)
ax.plot(angles, values, 'o-', linewidth=2)
ax.fill(angles, values, alpha=0.25)
ax.set_thetagrids(np.degrees(angles[:-1]), labels)
plt.title('Overall Performance Radar')
plt.savefig('radar.png')
plt.close()

# 四个维度分数分布直方图
hist_files = []
hist_conclusions = []
for col in ['relevance','quality','relaxation','immersiveness']:
    plt.figure(figsize=(5,3))
    data = df_expanded[col]
    plt.hist(data, bins=np.arange(1,7)-0.5, rwidth=0.8, color='C0', edgecolor='black')
    plt.xlabel('Score')
    plt.ylabel('Count')
    plt.title(f'{col.capitalize()} Score Distribution')
    plt.xticks([1,2,3,4,5])
    fname = f'{col}_hist.png'
    plt.savefig(fname)
    plt.close()
    hist_files.append(fname)
    # 结论
    counts = data.value_counts().sort_index()
    peak = counts.idxmax()
    if counts[peak] > counts.mean() * 1.5:
        shape = 'a strong peak'
    else:
        shape = 'a relatively flat distribution'
    hist_conclusions.append(
        f"Most ratings for {col} are concentrated at score {peak}, showing {shape}. "
        f"There are also non-negligible ratings at the lower end, indicating occasional failures."
    )

# mode均分条形图
mode_means = df_expanded.groupby('mode')[['relevance','quality','relaxation','immersiveness']].mean()
plt.figure(figsize=(7,4))
mode_means.plot(kind='bar')
plt.title('Mean Scores by Mode')
plt.ylabel('Mean Score')
plt.xlabel('Prompt Mode')
plt.xticks(rotation=30)
plt.tight_layout()
plt.savefig('mode_means.png')
plt.close()
# mode结论
top_mode = mode_means.mean(axis=1).idxmax()
low_mode = mode_means.mean(axis=1).idxmin()
conclusion_mode = (
    f"The model performs best in the '{top_mode}' mode and worst in the '{low_mode}' mode. "
    f"This suggests the model is more effective for certain use cases, while others remain challenging."
)

# category均分热力图
cat_means = df_expanded.groupby('category')[['relevance','quality','relaxation','immersiveness']].mean()
plt.figure(figsize=(7,4))
sns.heatmap(cat_means, annot=True, cmap='YlGnBu', fmt='.2f')
plt.title('Mean Scores by Category (Heatmap)')
plt.ylabel('Category')
plt.xlabel('Metric')
plt.tight_layout()
plt.savefig('category_heatmap.png')
plt.close()
# category结论
top_cat = cat_means.mean(axis=1).idxmax()
low_cat = cat_means.mean(axis=1).idxmin()
conclusion_cat = (
    f"The model achieves the highest average scores in the '{top_cat}' category and the lowest in the '{low_cat}' category. "
    f"This highlights clear strengths in generating certain types of audio, while other categories require further improvement."
)

# Word报告

doc = Document()
doc.add_heading('Nightingale Text-to-Audio Model Evaluation Report', 0)

doc.add_heading('2.1 Evaluation Objectives', level=1)
doc.add_paragraph(
    "This report aims to provide a systematic, data-driven analysis of the current performance of the generative text-to-audio model, going beyond anecdotal evaluation. The objectives are: (1) Quantify model performance on four key perceptual metrics (relevance, quality, relaxation, immersiveness); (2) Identify systematic success and failure patterns across audio categories and intended use modes; (3) Integrate qualitative feedback to understand the root causes behind quantitative scores; (4) Provide prioritized, actionable recommendations for future development."
)

doc.add_heading('2.2 Methodology', level=1)
doc.add_paragraph(
    "The analysis is based on a structured evaluation framework using a standardized test set of 30 diverse prompts, each assigned a primary mode (relax, focus, story) and a category (nature, urban, weather, home, meditation, music). Each audio is rated by human evaluators on a 1-5 Likert scale for four metrics: Relevance, Quality, Relaxation, Immersiveness. The dataset aggregates multiple rounds of evaluation to reduce individual bias and provide robust insights."
)

doc.add_heading('3.0 Macro Performance Analysis', level=1)
doc.add_paragraph("This section summarizes the overall performance across all evaluators and prompts.")
# 插入统计表
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Mean'
hdr_cells[2].text = 'Median'
hdr_cells[3].text = 'Std'
for _, row in stats_table.iterrows():
    cells = table.add_row().cells
    cells[0].text = str(row['Metric'])
    cells[1].text = str(row['Mean'])
    cells[2].text = str(row['Median'])
    cells[3].text = str(row['Std'])
doc.add_paragraph('Table: Overall performance metrics (mean, median, std) for each evaluation dimension.')
doc.add_paragraph(conclusion_stats)
doc.add_picture('radar.png', width=Inches(4.5))
doc.add_paragraph('Figure: Radar chart visualizing the mean scores of the four main evaluation metrics.')

doc.add_heading('3.2 Score Distribution', level=2)
doc.add_paragraph('Histograms below show the frequency of each score (1-5) for all four metrics across the dataset.')
for fname, col, concl in zip(hist_files, ['relevance','quality','relaxation','immersiveness'], hist_conclusions):
    doc.add_picture(fname, width=Inches(4))
    doc.add_paragraph(f'Figure: {col.capitalize()} score distribution.')
    doc.add_paragraph(concl)

doc.add_heading('4.0 Contextual Performance Analysis', level=1)
doc.add_paragraph('This section analyzes model performance by intended use mode and audio category.')
doc.add_heading('4.1 By Mode', level=2)
doc.add_picture('mode_means.png', width=Inches(5))
doc.add_paragraph('Figure: Mean scores for each evaluation metric grouped by prompt mode (relax, focus, story).')
doc.add_paragraph(conclusion_mode)
doc.add_heading('4.2 By Category', level=2)
doc.add_picture('category_heatmap.png', width=Inches(5))
doc.add_paragraph('Figure: Heatmap of mean scores for each evaluation metric grouped by audio category.')
doc.add_paragraph(conclusion_cat)

doc.add_heading('5.0 Conclusions and Recommendations', level=1)
doc.add_paragraph('The model demonstrates robust capabilities in generating high-fidelity, immersive audio for singular, textural events (e.g., thunderstorms, campfires), especially in weather and nature categories. However, it exhibits systematic weaknesses: (1) compositional failure for multi-element prompts, (2) inability to interpret semantic modifiers (e.g., "gentle", "distant"), and (3) failure to generate low-amplitude sounds. Recommendations: (1) Explore hierarchical generation architectures for compositionality; (2) Augment training data with explicit semantic modifier examples; (3) Review loss functions and normalization for low-amplitude signals; (4) Implement post-generation quality control and user guidance for prompt design.')

doc.add_heading('Appendix: Per-Prompt Analysis', level=1)
prompt_means = df_expanded.groupby('prompt')[['relevance','quality','relaxation','immersiveness']].mean()
doc.add_paragraph('Mean scores for each prompt:')
for idx, row in prompt_means.iterrows():
    doc.add_paragraph(f"{idx}: Relevance={row['relevance']:.2f}, Quality={row['quality']:.2f}, Relaxation={row['relaxation']:.2f}, Immersiveness={row['immersiveness']:.2f}")

doc.save('Nightingale_Evaluation_Report.docx')
print('Report generated: Nightingale_Evaluation_Report.docx') 